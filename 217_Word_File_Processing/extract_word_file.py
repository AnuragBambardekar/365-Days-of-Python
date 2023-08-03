from docx.api import Document

document = Document("test.docx")

# Print Headings and Titles
for p in document.paragraphs:
    if p.style.name.startswith("Heading") or p.style.name == "Title":
        print(p.text)

# Print the tables
for table in document.tables:
    print("Table: ")
    for row in table.rows:
        print(" | ".join([cell.text for cell in row.cells]))

# Print all the text in the document
all_text = ""
for p in document.paragraphs:
    all_text += p.text
    all_text += "\n"

print(all_text)

