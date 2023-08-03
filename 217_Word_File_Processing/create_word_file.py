# pip install python-docx
# pip install paragaphs

from docx import Document
import paragraphs

document = Document()

# Headings
document.add_heading("Hello World!", 0) # 0 is Title, 1,2,3,4 are Headings
document.add_heading("This is Heading 1", 1) # 0 is Title, 1,2,3,4 are Headings

# Paragraphs
p = document.add_paragraph("Sample Text. ")
p.add_run("This Text is BOLD. ").bold = True
p.add_run("This Text is italics. ").italic = True

# Bulleted list
document.add_paragraph("This is item one", style="List Bullet")
document.add_paragraph("This is item two", style="List Bullet")
document.add_paragraph("This is item three", style="List Bullet")
document.add_paragraph("This is item four", style="List Bullet")
document.add_paragraph("This is item five", style="List Bullet")

# Table
table_header = ["Name","Age","Job"]
some_data = [
    ["John",56,"Programmer"],
    ["Mary",55,"Accountant"],
    ["Anna",25,"Doctor"],
    ["Bob", 30,"Chef"]
]

table = document.add_table(rows=1, cols=3)
for i in range(3):
    table.rows[0].cells[i].text = table_header[i]

for name,age,job in some_data:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = str(age)
    cells[2].text = job

# Page Break
document.add_page_break()

# New Page
document.add_paragraph("This is a new page.")

# Add Image
document.add_picture("image.jpeg")

# Save the document
document.save("test.docx")